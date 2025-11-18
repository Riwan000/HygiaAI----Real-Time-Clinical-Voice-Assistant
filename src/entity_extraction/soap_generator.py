"""
SOAP Note Generator for Clinical Documentation

Converts raw consultation transcripts into structured SOAP (Subjective/Objective/Assessment/Plan) format.
Each SOAP field is embedded separately for better similarity recall in Qdrant.
Supports export to PDF and DOCX formats for clinical documentation.
"""

import logging
from typing import Dict, Any, Optional, List
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from ..entity_extraction.medical_ner import MedicalNER, MedicalEntity, EntityType

# Optional RAG enhancement
try:
    from ..rag.soap_rag_enhancer import SOAPRAGEnhancer
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False
    logger.warning("SOAP RAG enhancer not available. Enhanced extraction disabled.")

logger = logging.getLogger(__name__)

# Optional imports for document generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not available. PDF export will be disabled.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX export will be disabled.")


@dataclass
class SOAPNote:
    """Structured SOAP note with separate fields"""
    subjective: str
    objective: str
    assessment: str
    plan: str
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert SOAP note to dictionary"""
        return {
            "subjective": self.subjective,
            "objective": self.objective,
            "assessment": self.assessment,
            "plan": self.plan,
            "metadata": self.metadata
        }
    
    def get_field_embeddings_dict(self) -> Dict[str, str]:
        """Get dictionary of field names to text for embedding generation"""
        return {
            "subjective": self.subjective,
            "objective": self.objective,
            "assessment": self.assessment,
            "plan": self.plan
        }
    
    def export_to_pdf(self, output_path: str, patient_info: Optional[Dict[str, Any]] = None, 
                     clinician_info: Optional[Dict[str, Any]] = None) -> bool:
        """
        Export SOAP note to PDF format with professional clinical documentation layout
        
        Args:
            output_path: Path where PDF file should be saved
            patient_info: Optional patient information (name, ID, DOB, age, gender, etc.)
            clinician_info: Optional clinician information (name, title, license, etc.)
            
        Returns:
            True if export successful, False otherwise
        """
        if not REPORTLAB_AVAILABLE:
            logger.error("reportlab not available. Cannot export to PDF.")
            return False
        
        try:
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(output_path, pagesize=letter,
                                   rightMargin=0.75*inch, leftMargin=0.75*inch,
                                   topMargin=0.75*inch, bottomMargin=0.75*inch)
            story = []
            styles = getSampleStyleSheet()
            
            # Professional styles
            header_style = ParagraphStyle(
                'HeaderStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor='#000000',
                spaceAfter=8,
                alignment=TA_LEFT
            )
            
            title_style = ParagraphStyle(
                'TitleStyle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor='#1a1a1a',
                spaceAfter=16,
                spaceBefore=8,
                alignment=TA_LEFT,
                fontName='Helvetica-Bold'
            )
            
            section_title_style = ParagraphStyle(
                'SectionTitleStyle',
                parent=styles['Heading2'],
                fontSize=13,
                textColor='#2c3e50',
                spaceAfter=8,
                spaceBefore=16,
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                borderWidth=1,
                borderColor='#2c3e50',
                borderPadding=4
            )
            
            body_style = ParagraphStyle(
                'BodyTextStyle',
                parent=styles['Normal'],
                fontSize=11,
                textColor='#333333',
                spaceAfter=8,
                alignment=TA_LEFT,
                leading=14,
                leftIndent=0.2*inch
            )
            
            label_style = ParagraphStyle(
                'LabelStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor='#555555',
                spaceAfter=4,
                fontName='Helvetica-Bold'
            )
            
            # Professional Header Section
            header_data = []
            if clinician_info:
                header_data.append([
                    Paragraph(f"<b>Provider:</b> {clinician_info.get('name', 'N/A')}", header_style),
                    Paragraph(f"<b>Date:</b> {patient_info.get('date', datetime.now().strftime('%Y-%m-%d')) if patient_info else datetime.now().strftime('%Y-%m-%d')}", header_style)
                ])
                if clinician_info.get('title'):
                    header_data.append([
                        Paragraph(f"<b>Title:</b> {clinician_info.get('title')}", header_style),
                        Paragraph(f"<b>Time:</b> {datetime.now().strftime('%H:%M')}", header_style)
                    ])
            
            if header_data:
                header_table = Table(header_data, colWidths=[3.5*inch, 3.5*inch])
                header_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(header_table)
                story.append(Spacer(1, 0.2*inch))
            
            # Title
            story.append(Paragraph("SOAP NOTE", title_style))
            story.append(Spacer(1, 0.15*inch))
            
            # Patient Demographics Section (if provided)
            if patient_info:
                demo_data = []
                demo_row1 = []
                demo_row2 = []
                
                if patient_info.get('name'):
                    demo_row1.append(Paragraph(f"<b>Patient Name:</b> {patient_info.get('name')}", label_style))
                if patient_info.get('patient_id'):
                    demo_row1.append(Paragraph(f"<b>Patient ID:</b> {patient_info.get('patient_id')}", label_style))
                if demo_row1:
                    demo_data.append(demo_row1)
                
                if patient_info.get('dob') or patient_info.get('age'):
                    demo_row2.append(Paragraph(f"<b>DOB/Age:</b> {patient_info.get('dob', 'N/A')} / {patient_info.get('age', 'N/A')}", label_style))
                if patient_info.get('gender'):
                    demo_row2.append(Paragraph(f"<b>Gender:</b> {patient_info.get('gender')}", label_style))
                if demo_row2:
                    demo_data.append(demo_row2)
                
                if demo_data:
                    demo_table = Table(demo_data, colWidths=[3.5*inch, 3.5*inch])
                    demo_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f5f5f5')),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('LEFTPADDING', (0, 0), (-1, -1), 8),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(demo_table)
                    story.append(Spacer(1, 0.25*inch))
            
            # Subjective Section
            story.append(Paragraph("S - SUBJECTIVE", section_title_style))
            # Format subjective text with proper line breaks
            subjective_lines = self.subjective.split('\n')
            for line in subjective_lines:
                if line.strip():
                    if line.strip().startswith(('Chief Complaint:', 'History of Present Illness:', 
                                               'Medical History:', 'Current Medications:')):
                        # Make labels bold
                        parts = line.split(':', 1)
                        if len(parts) == 2:
                            story.append(Paragraph(f"<b>{parts[0]}:</b>{parts[1]}", body_style))
                        else:
                            story.append(Paragraph(line, body_style))
                    else:
                        story.append(Paragraph(line, body_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Objective Section
            story.append(Paragraph("O - OBJECTIVE", section_title_style))
            # Format objective text
            objective_lines = self.objective.split('\n')
            for line in objective_lines:
                if line.strip():
                    if line.strip().startswith(('Vital Signs:', 'Physical Examination:', 
                                               'Appearance and Behavior:', 'Lab/Test Results:')):
                        # Make section headers bold
                        parts = line.split(':', 1)
                        if len(parts) == 2:
                            story.append(Paragraph(f"<b>{parts[0]}:</b>{parts[1]}", body_style))
                        else:
                            story.append(Paragraph(line, body_style))
                    elif line.strip().startswith('  -'):
                        # Bullet points
                        story.append(Paragraph(f"• {line.strip()[2:].strip()}", body_style))
                    else:
                        story.append(Paragraph(line, body_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Assessment Section
            story.append(Paragraph("A - ASSESSMENT", section_title_style))
            assessment_lines = self.assessment.split('\n')
            for line in assessment_lines:
                if line.strip():
                    if line.strip().startswith(('Primary Diagnosis:', 'Clinical Impression:', 
                                               'Differential Diagnosis:', 'Clinical Reasoning:')):
                        parts = line.split(':', 1)
                        if len(parts) == 2:
                            story.append(Paragraph(f"<b>{parts[0]}:</b>{parts[1]}", body_style))
                        else:
                            story.append(Paragraph(line, body_style))
                    elif line.strip().startswith('  -'):
                        story.append(Paragraph(f"• {line.strip()[2:].strip()}", body_style))
                    else:
                        story.append(Paragraph(line, body_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Plan Section
            story.append(Paragraph("P - PLAN", section_title_style))
            plan_lines = self.plan.split('\n')
            for line in plan_lines:
                if line.strip():
                    if line.strip().startswith(('Medications:', 'Treatment Plan:', 
                                               'Diagnostic Tests/Orders:', 'Follow-up Instructions:',
                                               'Goals:', 'Patient Instructions/Education:')):
                        parts = line.split(':', 1)
                        if len(parts) == 2:
                            story.append(Paragraph(f"<b>{parts[0]}:</b>{parts[1]}", body_style))
                        else:
                            story.append(Paragraph(line, body_style))
                    elif line.strip().startswith('  -'):
                        story.append(Paragraph(f"• {line.strip()[2:].strip()}", body_style))
                    else:
                        story.append(Paragraph(line, body_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Signature Section
            story.append(Spacer(1, 0.2*inch))
            signature_data = [
                [Paragraph("_________________________", styles['Normal']), 
                 Paragraph("_________________________", styles['Normal'])],
                [Paragraph("<i>Clinician Signature</i>", styles['Italic']), 
                 Paragraph("<i>Date</i>", styles['Italic'])]
            ]
            signature_table = Table(signature_data, colWidths=[3.5*inch, 3.5*inch])
            signature_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(signature_table)
            
            # Footer with metadata (smaller, at bottom)
            if self.metadata:
                story.append(Spacer(1, 0.2*inch))
                footer_style = ParagraphStyle(
                    'FooterStyle',
                    parent=styles['Normal'],
                    fontSize=8,
                    textColor='#888888',
                    alignment=TA_LEFT
                )
                footer_text = f"Document Generated: {self.metadata.get('generated_at', 'N/A')[:10]} | "
                footer_text += f"Entity Count: {self.metadata.get('entity_count', 0)}"
                story.append(Paragraph(footer_text, footer_style))
            
            # Build PDF
            doc.build(story)
            logger.info(f"SOAP note exported to PDF: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting SOAP note to PDF: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def export_to_docx(self, output_path: str, patient_info: Optional[Dict[str, Any]] = None,
                      clinician_info: Optional[Dict[str, Any]] = None) -> bool:
        """
        Export SOAP note to DOCX format with professional clinical documentation layout
        
        Args:
            output_path: Path where DOCX file should be saved
            patient_info: Optional patient information (name, ID, DOB, age, gender, etc.)
            clinician_info: Optional clinician information (name, title, license, etc.)
            
        Returns:
            True if export successful, False otherwise
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available. Cannot export to DOCX.")
            return False
        
        try:
            from docx.shared import RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Professional Header Section
            if clinician_info:
                header_table = doc.add_table(rows=2, cols=2)
                header_table.style = 'Light Grid Accent 1'
                header_table.cell(0, 0).text = f"Provider: {clinician_info.get('name', 'N/A')}"
                header_table.cell(0, 1).text = f"Date: {patient_info.get('date', datetime.now().strftime('%Y-%m-%d')) if patient_info else datetime.now().strftime('%Y-%m-%d')}"
                if clinician_info.get('title'):
                    header_table.cell(1, 0).text = f"Title: {clinician_info.get('title')}"
                header_table.cell(1, 1).text = f"Time: {datetime.now().strftime('%H:%M')}"
                doc.add_paragraph()
            
            # Title
            title = doc.add_heading('SOAP NOTE', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_run = title.runs[0]
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            
            # Patient Demographics Table
            if patient_info:
                demo_table = doc.add_table(rows=2, cols=2)
                demo_table.style = 'Light Shading Accent 1'
                
                row_idx = 0
                if patient_info.get('name'):
                    demo_table.cell(row_idx, 0).text = f"Patient Name: {patient_info.get('name')}"
                if patient_info.get('patient_id'):
                    demo_table.cell(row_idx, 1).text = f"Patient ID: {patient_info.get('patient_id')}"
                
                row_idx = 1
                if patient_info.get('dob') or patient_info.get('age'):
                    dob_age = f"{patient_info.get('dob', 'N/A')} / {patient_info.get('age', 'N/A')}"
                    demo_table.cell(row_idx, 0).text = f"DOB/Age: {dob_age}"
                if patient_info.get('gender'):
                    demo_table.cell(row_idx, 1).text = f"Gender: {patient_info.get('gender')}"
                
                # Make labels bold in table cells
                for row in demo_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if ':' in paragraph.text:
                                parts = paragraph.text.split(':', 1)
                                paragraph.clear()
                                run1 = paragraph.add_run(parts[0] + ':')
                                run1.bold = True
                                if len(parts) > 1:
                                    paragraph.add_run(parts[1])
                
                doc.add_paragraph()
            
            # Subjective Section
            subjective_heading = doc.add_heading('S - SUBJECTIVE', level=1)
            subjective_heading.style.font.size = Pt(13)
            subjective_heading.runs[0].font.bold = True
            subjective_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            # Format subjective with proper structure
            subjective_lines = self.subjective.split('\n')
            for line in subjective_lines:
                if line.strip():
                    para = doc.add_paragraph()
                    para.style = 'List Paragraph'
                    para.paragraph_format.left_indent = Inches(0.2)
                    
                    if line.strip().startswith(('Chief Complaint:', 'History of Present Illness:',
                                               'Medical History:', 'Current Medications:',
                                               'Patient reports symptoms:')):
                        parts = line.split(':', 1)
                        if len(parts) == 2:
                            run1 = para.add_run(parts[0] + ':')
                            run1.bold = True
                            run1.font.size = Pt(11)
                            para.add_run(parts[1]).font.size = Pt(11)
                        else:
                            para.add_run(line).font.size = Pt(11)
                    else:
                        para.add_run(line.strip()).font.size = Pt(11)
            
            doc.add_paragraph()
            
            # Objective Section
            objective_heading = doc.add_heading('O - OBJECTIVE', level=1)
            objective_heading.style.font.size = Pt(13)
            objective_heading.runs[0].font.bold = True
            objective_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            objective_lines = self.objective.split('\n')
            for line in objective_lines:
                if line.strip():
                    para = doc.add_paragraph()
                    para.style = 'List Paragraph'
                    para.paragraph_format.left_indent = Inches(0.2)
                    
                    if line.strip().startswith(('Vital Signs:', 'Physical Examination:',
                                               'Appearance and Behavior:', 'Lab/Test Results:',
                                               'Procedures Performed:', 'Lab Tests Ordered:')):
                        parts = line.split(':', 1)
                        if len(parts) == 2:
                            run1 = para.add_run(parts[0] + ':')
                            run1.bold = True
                            run1.font.size = Pt(11)
                            para.add_run(parts[1]).font.size = Pt(11)
                        else:
                            para.add_run(line).font.size = Pt(11)
                    elif line.strip().startswith('  -'):
                        para.add_run('• ' + line.strip()[2:].strip()).font.size = Pt(11)
                    else:
                        para.add_run(line.strip()).font.size = Pt(11)
            
            doc.add_paragraph()
            
            # Assessment Section
            assessment_heading = doc.add_heading('A - ASSESSMENT', level=1)
            assessment_heading.style.font.size = Pt(13)
            assessment_heading.runs[0].font.bold = True
            assessment_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            assessment_lines = self.assessment.split('\n')
            for line in assessment_lines:
                if line.strip():
                    para = doc.add_paragraph()
                    para.style = 'List Paragraph'
                    para.paragraph_format.left_indent = Inches(0.2)
                    
                    if line.strip().startswith(('Primary Diagnosis:', 'Clinical Impression:',
                                               'Differential Diagnosis:', 'Clinical Reasoning:')):
                        parts = line.split(':', 1)
                        if len(parts) == 2:
                            run1 = para.add_run(parts[0] + ':')
                            run1.bold = True
                            run1.font.size = Pt(11)
                            para.add_run(parts[1]).font.size = Pt(11)
                        else:
                            para.add_run(line).font.size = Pt(11)
                    elif line.strip().startswith('  -'):
                        para.add_run('• ' + line.strip()[2:].strip()).font.size = Pt(11)
                    else:
                        para.add_run(line.strip()).font.size = Pt(11)
            
            doc.add_paragraph()
            
            # Plan Section
            plan_heading = doc.add_heading('P - PLAN', level=1)
            plan_heading.style.font.size = Pt(13)
            plan_heading.runs[0].font.bold = True
            plan_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            plan_lines = self.plan.split('\n')
            for line in plan_lines:
                if line.strip():
                    para = doc.add_paragraph()
                    para.style = 'List Paragraph'
                    para.paragraph_format.left_indent = Inches(0.2)
                    
                    if line.strip().startswith(('Medications:', 'Treatment Plan:',
                                               'Diagnostic Tests/Orders:', 'Follow-up Instructions:',
                                               'Goals:', 'Patient Instructions/Education:')):
                        parts = line.split(':', 1)
                        if len(parts) == 2:
                            run1 = para.add_run(parts[0] + ':')
                            run1.bold = True
                            run1.font.size = Pt(11)
                            para.add_run(parts[1]).font.size = Pt(11)
                        else:
                            para.add_run(line).font.size = Pt(11)
                    elif line.strip().startswith('  -'):
                        para.add_run('• ' + line.strip()[2:].strip()).font.size = Pt(11)
                    else:
                        para.add_run(line.strip()).font.size = Pt(11)
            
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Signature Section
            sig_table = doc.add_table(rows=2, cols=2)
            sig_table.cell(0, 0).text = "_________________________"
            sig_table.cell(0, 1).text = "_________________________"
            sig_para1 = sig_table.cell(1, 0).paragraphs[0]
            sig_para1.add_run("Clinician Signature").italic = True
            sig_para1.runs[0].font.size = Pt(10)
            sig_para2 = sig_table.cell(1, 1).paragraphs[0]
            sig_para2.add_run("Date").italic = True
            sig_para2.runs[0].font.size = Pt(10)
            
            # Footer with metadata
            if self.metadata:
                doc.add_paragraph()
                footer = doc.add_paragraph()
                footer_run1 = footer.add_run(f"Document Generated: {self.metadata.get('generated_at', 'N/A')[:10]} | ")
                footer_run1.italic = True
                footer_run1.font.size = Pt(8)
                footer_run1.font.color.rgb = RGBColor(136, 136, 136)
                footer_run2 = footer.add_run(f"Entity Count: {self.metadata.get('entity_count', 0)}")
                footer_run2.italic = True
                footer_run2.font.size = Pt(8)
                footer_run2.font.color.rgb = RGBColor(136, 136, 136)
            
            # Save document
            doc.save(output_path)
            logger.info(f"SOAP note exported to DOCX: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting SOAP note to DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False


class SOAPGenerator:
    """
    Generates structured SOAP notes from consultation transcripts
    
    Pipeline: Transcript → Entity Extraction → SOAP Classification → Structured Output
    """
    
    def __init__(
        self,
        ner_model: Optional[MedicalNER] = None,
        use_rag: bool = True
    ):
        """
        Initialize SOAP generator
        
        Args:
            ner_model: Optional MedicalNER instance (creates new if not provided)
            use_rag: Whether to use RAG enhancement from knowledge base (default: True)
        """
        self.ner = ner_model or MedicalNER()
        self.use_rag = use_rag and RAG_AVAILABLE
        
        # Initialize RAG enhancer if available and enabled
        self.rag_enhancer = None
        if self.use_rag:
            try:
                self.rag_enhancer = SOAPRAGEnhancer()
                logger.info("SOAP generator initialized with RAG enhancement")
            except Exception as e:
                logger.warning(f"Failed to initialize RAG enhancer: {e}. Continuing without RAG.")
                self.use_rag = False
        else:
            logger.info("SOAP generator initialized (RAG disabled)")
    
    def generate_soap(
        self,
        transcript: str,
        entities: Optional[List[MedicalEntity]] = None,
        patient_metadata: Optional[Dict[str, Any]] = None
    ) -> SOAPNote:
        """
        Generate structured SOAP note from transcript using LLM and knowledge base
        
        Args:
            transcript: Raw consultation transcript text
            entities: Optional pre-extracted entities (will extract if not provided)
            patient_metadata: Optional patient metadata (age, gender, etc.)
            
        Returns:
            SOAPNote with structured S/O/A/P sections
        """
        # Extract entities if not provided
        if entities is None:
            entities = self.ner.extract_entities(transcript)
        
        # Retrieve relevant knowledge base information for RAG enhancement
        kb_context = self._retrieve_knowledge_base_context(transcript, entities, patient_metadata)
        
        # Use LLM to generate structured SOAP note
        soap_note = self._generate_soap_with_llm(transcript, entities, kb_context, patient_metadata)
        
        # Build metadata
        metadata = {
            "generated_at": datetime.utcnow().isoformat(),
            "entity_count": len(entities),
            "transcript_length": len(transcript),
            "patient_metadata": patient_metadata or {},
            "llm_generated": True,
            "rag_enhanced": bool(kb_context)
        }
        
        return SOAPNote(
            subjective=soap_note["subjective"],
            objective=soap_note["objective"],
            assessment=soap_note["assessment"],
            plan=soap_note["plan"],
            metadata=metadata
        )
    
    def _retrieve_knowledge_base_context(
        self,
        transcript: str,
        entities: List[MedicalEntity],
        patient_metadata: Optional[Dict[str, Any]]
    ) -> str:
        """Retrieve relevant knowledge base context for SOAP generation"""
        # Always try to retrieve knowledge base context for better SOAP structure
        try:
            import os
            from src.storage.qdrant_storage import QdrantStorage
            from src.embeddings import BioBERTEmbeddingGenerator
            
            # Use clinical_kb_collection for knowledge base
            qdrant_url = os.getenv("QDRANT_URL")
            if qdrant_url:
                kb_storage = QdrantStorage(
                    url=qdrant_url,
                    api_key=os.getenv("QDRANT_API_KEY"),
                    collection_name="clinical_kb_collection",
                    vector_size=768
                )
            else:
                kb_storage = QdrantStorage(
                    host=os.getenv("QDRANT_HOST", "localhost"),
                    port=int(os.getenv("QDRANT_PORT", "6334")),
                    api_key=os.getenv("QDRANT_API_KEY"),
                    collection_name="clinical_kb_collection",
                    vector_size=768
                )
            
            embedder = BioBERTEmbeddingGenerator()
            
            # Build query from transcript and entities
            symptoms = [e.text for e in entities if e.entity_type == EntityType.SYMPTOM]
            diagnoses = [e.text for e in entities if e.entity_type == EntityType.DIAGNOSIS]
            diseases = [e.text for e in entities if e.entity_type == EntityType.DISEASE]
            
            # Retrieve SOAP note writing guidelines
            soap_query = "SOAP note writing guide structure format subjective objective assessment plan"
            soap_embedding = embedder.generate_embedding(soap_query)
            soap_results = kb_storage.search_with_filters(
                query_embedding=soap_embedding,
                filters={"domain": "guidelines"},
                limit=3,
                score_threshold=0.3
            )
            
            # Retrieve disease-specific information if diagnoses found
            disease_results = []
            if diagnoses or diseases:
                disease_query = f"{' '.join((diagnoses + diseases)[:2])} clinical information diagnosis treatment"
                disease_embedding = embedder.generate_embedding(disease_query)
                disease_results = kb_storage.search_with_filters(
                    query_embedding=disease_embedding,
                    limit=3,
                    score_threshold=0.3
                )
            
            # Format knowledge base context
            context_parts = []
            
            # Add SOAP guidelines
            for result in soap_results:
                payload = result.get("payload", {})
                content = payload.get("content", "") or payload.get("text", "")
                title = payload.get("title", "")
                if content:
                    context_parts.append(f"SOAP Guidelines ({title}):\n{content[:500]}")
            
            # Add disease information
            for result in disease_results:
                payload = result.get("payload", {})
                content = payload.get("content", "") or payload.get("text", "")
                title = payload.get("title", "")
                if content:
                    context_parts.append(f"Disease Information ({title}):\n{content[:500]}")
            
            return "\n\n".join(context_parts)
            
        except Exception as e:
            logger.warning(f"Error retrieving knowledge base context: {e}")
            return ""
    
    def _generate_soap_with_llm(
        self,
        transcript: str,
        entities: List[MedicalEntity],
        kb_context: str,
        patient_metadata: Optional[Dict[str, Any]]
    ) -> Dict[str, str]:
        """Generate SOAP note using Gemini LLM with knowledge base context"""
        try:
            import os
            import google.generativeai as genai
            import json
            
            # Initialize Gemini
            api_key = os.getenv("GOOGLE_API_KEY")
            if not api_key:
                logger.warning("GOOGLE_API_KEY not found. Falling back to rule-based extraction.")
                return self._generate_soap_rule_based(transcript, entities)
            
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.5-flash")
            
            # Add patient metadata if available
            metadata_text = ""
            if patient_metadata:
                metadata_parts = []
                if patient_metadata.get("age_group"):
                    metadata_parts.append(f"Age Group: {patient_metadata['age_group']}")
                if patient_metadata.get("region"):
                    metadata_parts.append(f"Region: {patient_metadata['region']}")
                if patient_metadata.get("comorbidities"):
                    comorbidities = patient_metadata["comorbidities"]
                    if isinstance(comorbidities, list):
                        metadata_parts.append(f"Comorbidities: {', '.join(comorbidities)}")
                    else:
                        metadata_parts.append(f"Comorbidities: {comorbidities}")
                if metadata_parts:
                    metadata_text = "\n".join(metadata_parts)
            
            # Extract key entities for context
            symptoms = [e.text for e in entities if e.entity_type == EntityType.SYMPTOM]
            diagnoses = [e.text for e in entities if e.entity_type == EntityType.DIAGNOSIS]
            medications = [e.text for e in entities if e.entity_type == EntityType.MEDICATION]
            
            # Build main prompt with explicit extraction instructions
            prompt = f"""You are a medical professional generating a structured SOAP note from a clinical consultation transcript.

PATIENT INFORMATION:
{metadata_text if metadata_text else "Not provided"}

EXTRACTED MEDICAL ENTITIES:
- Symptoms: {', '.join(symptoms) if symptoms else 'None identified'}
- Diagnoses: {', '.join(diagnoses) if diagnoses else 'None identified'}
- Medications: {', '.join(medications) if medications else 'None identified'}

CONSULTATION TRANSCRIPT:
{transcript[:3000] if len(transcript) > 3000 else transcript}

CRITICAL INSTRUCTIONS:
1. DO NOT copy the entire transcript verbatim into any section
2. EXTRACT and SUMMARIZE only relevant information for each section
3. Each section should be concise and contain only extracted, relevant information
4. Use the knowledge base context to ensure clinical accuracy and proper structure
5. Format each section with clear subsections using line breaks

REQUIRED OUTPUT FORMAT:

SUBJECTIVE (S) - Extract ONLY patient-reported information:
- Chief Complaint: [Brief statement of main concern - 1-2 sentences]
- History of Present Illness: [Key details about current symptoms, when they started, duration, severity - 3-5 sentences]
- Medical History: [Relevant past conditions mentioned - bullet points]
- Current Medications: [List medications mentioned - bullet points]
- Social History: [Relevant social information if mentioned]

OBJECTIVE (O) - Extract ONLY observable/measured findings:
- Vital Signs: [Only if mentioned - e.g., "BP 120/80, HR 72, Temp 98.6F" or "Not documented"]
- Physical Examination: [Only objective findings mentioned - bullet points or "Not documented"]
- Appearance and Behavior: [Only if mentioned - e.g., "Alert, oriented, cooperative" or "Not documented"]
- Lab/Test Results: [Only if mentioned - bullet points or "Not documented"]

ASSESSMENT (A) - Extract clinical reasoning and diagnoses:
- Primary Diagnosis: [Main diagnosis mentioned or inferred - 1-2 sentences]
- Clinical Impression: [Clinical reasoning based on findings - 2-4 sentences]
- Differential Diagnosis: [Other possibilities mentioned - bullet points or "None discussed"]
- Clinical Reasoning: [Why this diagnosis fits - 2-3 sentences]

PLAN (P) - Extract treatment plans and recommendations:
- Medications: [Prescribed medications with dosages if mentioned - bullet points or "No new medications"]
- Treatment Plan: [Treatment recommendations discussed - bullet points]
- Diagnostic Tests/Orders: [Tests ordered or recommended - bullet points or "None"]
- Follow-up Instructions: [When to return, monitoring - bullet points or "Not specified"]
- Patient Instructions/Education: [Patient education provided - bullet points or "None"]

CRITICAL RULES:
- Each section must be DISTINCT and contain DIFFERENT information
- DO NOT repeat the same information across sections
- DO NOT include the full transcript in any section
- Each section should be 50-300 words maximum
- Use bullet points for lists
- Use professional medical language
- If information is not present, write "Not documented" or "Not mentioned"

Generate ONLY a JSON object with this exact structure:
{{
    "subjective": "Chief Complaint: ...\\nHistory of Present Illness: ...\\nMedical History: ...\\nCurrent Medications: ...",
    "objective": "Vital Signs: ...\\nPhysical Examination: ...\\nAppearance and Behavior: ...\\nLab/Test Results: ...",
    "assessment": "Primary Diagnosis: ...\\nClinical Impression: ...\\nDifferential Diagnosis: ...\\nClinical Reasoning: ...",
    "plan": "Medications: ...\\nTreatment Plan: ...\\nDiagnostic Tests/Orders: ...\\nFollow-up Instructions: ...\\nPatient Instructions/Education: ..."
}}

Return ONLY valid JSON, no additional text before or after.
"""
            
            # Combine prompt parts if KB context exists
            full_prompt = prompt
            if kb_context:
                full_prompt = f"""MEDICAL KNOWLEDGE BASE CONTEXT:
{kb_context}

Use the SOAP note writing guidelines above to structure your output properly.

{prompt}"""
            
            # Generate response
            response = model.generate_content(
                full_prompt,
                generation_config={
                    "temperature": 0.2,  # Lower temperature for more structured output
                    "top_p": 0.95,
                    "top_k": 40,
                    "max_output_tokens": 3000,  # Increased for better extraction
                }
            )
            
            # Parse JSON response
            response_text = response.text.strip()
            
            # Try to extract JSON from response (may have markdown code blocks)
            if "```json" in response_text:
                json_start = response_text.find("```json") + 7
                json_end = response_text.find("```", json_start)
                response_text = response_text[json_start:json_end].strip()
            elif "```" in response_text:
                json_start = response_text.find("```") + 3
                json_end = response_text.find("```", json_start)
                response_text = response_text[json_start:json_end].strip()
            
            # Try to find JSON object boundaries if not in code blocks
            if "{" in response_text and "}" in response_text:
                json_start = response_text.find("{")
                json_end = response_text.rfind("}") + 1
                response_text = response_text[json_start:json_end]
            
            # Parse JSON
            soap_data = json.loads(response_text)
            
            # Validate that sections don't contain the full transcript
            transcript_length = len(transcript)
            subjective = soap_data.get("subjective", "").strip()
            objective = soap_data.get("objective", "").strip()
            assessment = soap_data.get("assessment", "").strip()
            plan = soap_data.get("plan", "").strip()
            
            # Check if any section is suspiciously long (likely contains full transcript)
            max_section_length = min(transcript_length * 0.8, 2000)  # Max 80% of transcript or 2000 chars
            
            if len(subjective) > max_section_length:
                logger.warning(f"Subjective section too long ({len(subjective)} chars), likely contains full transcript. Using rule-based fallback.")
                return self._generate_soap_rule_based(transcript, entities)
            if len(objective) > max_section_length:
                logger.warning(f"Objective section too long ({len(objective)} chars), likely contains full transcript. Using rule-based fallback.")
                return self._generate_soap_rule_based(transcript, entities)
            if len(assessment) > max_section_length:
                logger.warning(f"Assessment section too long ({len(assessment)} chars), likely contains full transcript. Using rule-based fallback.")
                return self._generate_soap_rule_based(transcript, entities)
            if len(plan) > max_section_length:
                logger.warning(f"Plan section too long ({len(plan)} chars), likely contains full transcript. Using rule-based fallback.")
                return self._generate_soap_rule_based(transcript, entities)
            
            # Check if sections are too similar (all contain same content)
            if subjective == objective or subjective == assessment or subjective == plan:
                logger.warning("SOAP sections are identical, likely contains full transcript. Using rule-based fallback.")
                return self._generate_soap_rule_based(transcript, entities)
            
            # Validate and return
            return {
                "subjective": subjective if subjective else "No subjective information recorded.",
                "objective": objective if objective else "No objective findings recorded.",
                "assessment": assessment if assessment else "Assessment pending further evaluation.",
                "plan": plan if plan else "Plan pending assessment completion."
            }
            
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing LLM JSON response: {e}")
            logger.error(f"Response text (first 500 chars): {response_text[:500]}")
            logger.info("Falling back to rule-based extraction")
            import traceback
            logger.debug(traceback.format_exc())
            return self._generate_soap_rule_based(transcript, entities)
        except Exception as e:
            logger.error(f"Error generating SOAP with LLM: {e}")
            logger.error(f"Response text (first 500 chars): {response_text[:500] if 'response_text' in locals() else 'N/A'}")
            logger.info("Falling back to rule-based extraction")
            import traceback
            logger.debug(traceback.format_exc())
            return self._generate_soap_rule_based(transcript, entities)
    
    def _generate_soap_rule_based(
        self,
        transcript: str,
        entities: List[MedicalEntity]
    ) -> Dict[str, str]:
        """Fallback rule-based SOAP generation"""
        subjective = self._extract_subjective(transcript, entities)
        objective = self._extract_objective(transcript, entities)
        assessment = self._generate_assessment(transcript, entities, subjective, objective)
        plan = self._generate_plan(transcript, entities, assessment)
        
        return {
            "subjective": subjective,
            "objective": objective,
            "assessment": assessment,
            "plan": plan
        }
    
    def _extract_subjective(self, transcript: str, entities: List[MedicalEntity]) -> str:
        """
        Extract Subjective section (patient-reported symptoms, history)
        
        Subjective includes (per SOAP best practices):
        - Chief Complaint (CC) or presenting problem
        - History of present illness
        - Medical history
        - Review of systems
        - Current medications
        - What patient/client says (attributed to them)
        - Information from family members
        - Past medical records review
        """
        # Enhance with knowledge base if RAG is enabled
        rag_context = None
        if self.use_rag and self.rag_enhancer:
            try:
                enhancement = self.rag_enhancer.enhance_subjective_extraction(transcript, entities)
                rag_context = self.rag_enhancer.format_knowledge_context(
                    enhancement.get("guidelines", []) + enhancement.get("symptom_knowledge", [])
                )
            except Exception as e:
                logger.warning(f"RAG enhancement failed: {e}. Continuing without enhancement.")
        
        subjective_parts = []
        
        # Extract Chief Complaint (CC) - usually first thing mentioned
        # Use knowledge base guidelines if available
        if rag_context and 'Chief Complaint' in rag_context:
            # Knowledge base provides structure guidance
            pass
        
        cc_keywords = ['chief complaint', 'presents with', 'complains of', 'complains about', 
                      'main concern', 'reason for visit', 'presenting problem', 'here for', 'came in for']
        lines = transcript.split('\n')
        chief_complaint = None
        
        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line_lower = line.lower().strip()
            for keyword in cc_keywords:
                if keyword in line_lower:
                    # Extract the sentence containing the keyword
                    sentences = line.split('.')
                    for sentence in sentences:
                        if keyword in sentence.lower():
                            chief_complaint = sentence.strip()
                            break
                    if chief_complaint:
                        break
            if chief_complaint:
                break
        
        # If no explicit CC found, look for first person statements or patient-reported content
        if not chief_complaint:
            for line in lines[:5]:
                line_lower = line.lower()
                # Look for patient-reported language
                if any(phrase in line_lower for phrase in ['patient reports', 'patient says', 
                      'patient states', 'reports feeling', 'reports having', 'i feel', 'i have']):
                    chief_complaint = line.strip()
                    break
        
        if chief_complaint:
            subjective_parts.append(f"Chief Complaint: {chief_complaint}")
        
        # Extract History of Present Illness (HPI)
        hpi_keywords = ['history of present illness', 'hpi', 'symptoms started', 'when did', 
                       'how long', 'duration', 'onset']
        hpi_sentences = []
        for line in lines:
            line_lower = line.lower()
            for keyword in hpi_keywords:
                if keyword in line_lower:
                    hpi_sentences.append(line.strip())
                    break
        
        # Also look for temporal information (when symptoms started)
        for line in lines:
            line_lower = line.lower()
            if any(phrase in line_lower for phrase in ['for the past', 'for', 'days', 'weeks', 'months', 
                  'since', 'started', 'began']):
                if 'patient' in line_lower or 'reports' in line_lower:
                    if line.strip() not in hpi_sentences:
                        hpi_sentences.append(line.strip())
        
        if hpi_sentences:
            subjective_parts.append(f"History of Present Illness: {' '.join(hpi_sentences[:3])}")
        
        # Extract patient-reported symptoms
        symptoms = [e.text for e in entities if e.entity_type == EntityType.SYMPTOM]
        if symptoms:
            # Use knowledge base to enhance symptom description
            symptom_text = ', '.join(symptoms)
            if rag_context and any(symptom.lower() in rag_context.lower() for symptom in symptoms[:2]):
                # Knowledge base has information about these symptoms
                subjective_parts.append(f"Patient reports symptoms: {symptom_text}")
            else:
                subjective_parts.append(f"Patient reports symptoms: {symptom_text}")
        
        # Extract medical history (conditions, past diagnoses)
        conditions = [e.text for e in entities if e.entity_type == EntityType.CONDITION]
        diseases = [e.text for e in entities if e.entity_type == EntityType.DISEASE]
        diagnoses = [e.text for e in entities if e.entity_type == EntityType.DIAGNOSIS]
        
        medical_history = []
        for line in lines:
            line_lower = line.lower()
            if any(phrase in line_lower for phrase in ['past medical history', 'medical history', 
                  'history of', 'previous', 'prior diagnosis', 'has a history']):
                medical_history.append(line.strip())
        
        if medical_history:
            subjective_parts.append(f"Medical History: {' '.join(medical_history[:2])}")
        elif conditions or diseases or diagnoses:
            all_conditions = list(set(conditions + diseases + diagnoses))
            subjective_parts.append(f"Past Medical History: {', '.join(all_conditions[:5])}")
        
        # Extract current medications
        medications = [e.text for e in entities if e.entity_type == EntityType.MEDICATION]
        if medications:
            subjective_parts.append(f"Current Medications: {', '.join(medications)}")
        
        # Add any other patient-reported information (first person, quotes, etc.)
        patient_reported = []
        for line in lines:
            line_lower = line.lower()
            if any(phrase in line_lower for phrase in ['patient reports', 'patient says', 
                  'patient states', 'reports that', 'states that']):
                if line.strip() not in subjective_parts:
                    patient_reported.append(line.strip())
        
        if patient_reported:
            subjective_parts.extend(patient_reported[:2])
        
        # If no structured extraction, try to extract intelligently
        if not subjective_parts:
            # Look for patient-reported statements more carefully
            sentences = transcript.split('.')
            patient_statements = []
            
            for sentence in sentences[:10]:  # Check first 10 sentences
                sentence_lower = sentence.lower().strip()
                # Look for patient-reported patterns
                if any(phrase in sentence_lower for phrase in [
                    'i have', 'i feel', 'i am', 'i\'m', 'my', 'patient reports',
                    'patient says', 'patient states', 'complains of', 'presents with'
                ]):
                    if len(sentence.strip()) > 20:  # Meaningful sentence
                        patient_statements.append(sentence.strip())
            
            if patient_statements:
                subjective_parts.append(f"Chief Complaint: {patient_statements[0]}")
                if len(patient_statements) > 1:
                    subjective_parts.append(f"History of Present Illness: {' '.join(patient_statements[1:3])}")
            else:
                # Last resort: take first meaningful portion
                first_portion = transcript[:int(len(transcript) * 0.2)].strip()
                if first_portion:
                    subjective_parts.append(f"Patient reports: {first_portion}")
        
        return "\n".join(subjective_parts) if subjective_parts else "No subjective information recorded."
    
    def _extract_objective(self, transcript: str, entities: List[MedicalEntity]) -> str:
        """
        Extract Objective section (observable findings, vital signs, exam findings)
        
        Objective includes (per SOAP best practices):
        - Vital signs (factual measurements)
        - Physical examination findings (what clinician observes)
        - Client/patient appearance, behavior, and mood in session
        - Relevant medical records from other specialists
        - Lab results (if available)
        - ONLY factual information observed, NOT what patient told you
        """
        # Enhance with knowledge base if RAG is enabled
        rag_context = None
        if self.use_rag and self.rag_enhancer:
            try:
                enhancement = self.rag_enhancer.enhance_objective_extraction(transcript, entities)
                rag_context = self.rag_enhancer.format_knowledge_context(
                    enhancement.get("vital_knowledge", []) + enhancement.get("exam_knowledge", [])
                )
            except Exception as e:
                logger.warning(f"RAG enhancement failed: {e}. Continuing without enhancement.")
        
        objective_parts = []
        lines = transcript.split('\n')
        
        # Extract vital signs (factual measurements)
        vitals = [e.text for e in entities if e.entity_type == EntityType.VITAL_SIGN]
        vital_signs_text = []
        
        # Look for vital signs with measurements
        vital_keywords = ['temperature', 'temp', 'blood pressure', 'bp', 'heart rate', 'hr', 
                         'pulse', 'respiratory rate', 'rr', 'oxygen saturation', 'o2 sat', 
                         'spo2', 'weight', 'height', 'bmi', 'blood pressure is', 'heart rate is',
                         'temperature is', 'vital signs']
        
        # Use knowledge base for normal ranges if available
        normal_ranges = {}
        if rag_context and 'Normal Vital Signs' in rag_context:
            # Extract normal ranges from knowledge base
            pass
        
        for line in lines:
            line_lower = line.lower()
            # Look for lines with vital sign keywords and numbers/measurements
            # Exclude patient-reported vital signs
            if any(keyword in line_lower for keyword in vital_keywords):
                # Make sure it's not patient-reported
                if 'patient reports' not in line_lower and 'patient says' not in line_lower:
                    # Check if line contains measurements (numbers)
                    if any(char.isdigit() for char in line):
                        vital_signs_text.append(line.strip())
        
        if vital_signs_text:
            objective_parts.append("Vital Signs:")
            objective_parts.extend([f"  - {v}" for v in vital_signs_text[:8]])  # Limit to 8 vital signs
        elif vitals:
            objective_parts.append(f"Vital Signs: {', '.join(vitals)}")
        
        # Extract physical examination findings (what clinician observes)
        exam_keywords = ['physical examination', 'examination reveals', 'exam shows', 
                        'inspection', 'palpation', 'auscultation', 'percussion', 
                        'observed', 'findings', 'appears', 'shows', 'demonstrates',
                        'examination of', 'chest examination', 'abdominal examination']
        
        exam_findings = []
        for line in lines:
            line_lower = line.lower()
            # Look for examination language (clinician observations)
            if any(keyword in line_lower for keyword in exam_keywords):
                # Make sure it's not patient-reported (avoid "patient reports examination")
                if 'patient reports' not in line_lower and 'patient says' not in line_lower:
                    exam_findings.append(line.strip())
        
        if exam_findings:
            objective_parts.append("\nPhysical Examination:")
            objective_parts.extend([f"  - {f}" for f in exam_findings[:5]])  # Limit to 5 findings
        
        # Extract appearance, behavior, and mood (clinician observations)
        appearance_keywords = ['appears', 'appearance', 'well-appearing', 'ill-appearing',
                              'alert', 'oriented', 'cooperative', 'uncooperative', 'anxious',
                              'calm', 'distressed', 'comfortable', 'uncomfortable']
        
        behavior_keywords = ['behavior', 'fidgety', 'wringing', 'speaking', 'concentrating',
                            'difficulty', 'mood', 'affect', 'flat', 'appropriate', 'inappropriate']
        
        observations = []
        for line in lines:
            line_lower = line.lower()
            # Look for clinician observations of appearance/behavior
            if any(keyword in line_lower for keyword in appearance_keywords + behavior_keywords):
                # Make sure it's an observation, not patient-reported
                if 'patient reports' not in line_lower and 'patient says' not in line_lower:
                    if any(obs_word in line_lower for obs_word in ['appeared', 'was', 'were', 'showed', 'demonstrated']):
                        observations.append(line.strip())
        
        if observations:
            objective_parts.append("\nAppearance and Behavior:")
            objective_parts.extend([f"  - {o}" for o in observations[:3]])
        
        # Extract procedures performed (objective)
        procedures = [e.text for e in entities if e.entity_type == EntityType.PROCEDURE]
        if procedures:
            objective_parts.append(f"\nProcedures Performed: {', '.join(procedures)}")
        
        # Extract lab tests/results (objective data)
        lab_tests = [e.text for e in entities if e.entity_type == EntityType.LAB_TEST]
        lab_results = []
        for line in lines:
            line_lower = line.lower()
            if any(phrase in line_lower for phrase in ['lab results', 'test results', 'cbc', 
                  'complete blood count', 'x-ray', 'chest x-ray', 'ct scan', 'mri']):
                lab_results.append(line.strip())
        
        if lab_results:
            objective_parts.append("\nLab/Test Results:")
            objective_parts.extend([f"  - {l}" for l in lab_results[:3]])
        elif lab_tests:
            objective_parts.append(f"\nLab Tests Ordered: {', '.join(lab_tests)}")
        
        # If no structured extraction, look for objective language more intelligently
        if not objective_parts or len(objective_parts) < 2:
            # Look for sentences with objective measurement language
            sentences = transcript.split('.')
            objective_sentences = []
            
            # Objective language patterns
            objective_patterns = [
                'measured', 'found', 'observed', 'reveals', 'shows', 'demonstrates',
                'on examination', 'physical exam', 'appears', 'is', 'was', 'were',
                'blood pressure', 'heart rate', 'temperature', 'vital signs',
                'alert', 'oriented', 'cooperative', 'well-appearing'
            ]
            
            for sentence in sentences:
                sentence_lower = sentence.lower().strip()
                # Look for objective measurement language
                if any(keyword in sentence_lower for keyword in objective_patterns):
                    # Exclude patient-reported
                    if all(phrase not in sentence_lower for phrase in [
                        'patient reports', 'patient says', 'patient states', 'i feel', 'i have'
                    ]):
                        # Check if it contains measurements or observations
                        if any(char.isdigit() for char in sentence) or len(sentence.strip()) > 30:
                            objective_sentences.append(sentence.strip())
            
            if objective_sentences:
                objective_parts.append("\nAdditional Observations:")
                objective_parts.extend([f"  - {s}" for s in objective_sentences[:5]])
            else:
                # Last resort: look for clinician statements
                clinician_statements = []
                for sentence in sentences:
                    sentence_lower = sentence.lower()
                    if any(phrase in sentence_lower for phrase in [
                        'doctor', 'physician', 'clinician', 'i would', 'i will', 'i recommend'
                    ]):
                        if 'patient reports' not in sentence_lower:
                            clinician_statements.append(sentence.strip())
                
                if clinician_statements:
                    objective_parts.append("\nClinician Observations:")
                    objective_parts.extend([f"  - {s}" for s in clinician_statements[:3]])
        
        return "\n".join(objective_parts) if objective_parts else "No objective findings recorded."
    
    def _generate_assessment(
        self,
        transcript: str,
        entities: List[MedicalEntity],
        subjective: str,
        objective: str
    ) -> str:
        """
        Generate Assessment section (diagnosis, differential diagnosis, clinical reasoning)
        
        Assessment includes (per SOAP best practices):
        - Combines information from Subjective and Objective sections
        - Clinician's impressions and interpretation
        - Diagnosis or list of possible diagnoses
        - Use of clinical knowledge/DSM-5 criteria/therapeutic models
        - Clinical reasoning
        """
        # Enhance with knowledge base if RAG is enabled
        rag_context = None
        if self.use_rag and self.rag_enhancer:
            try:
                enhancement = self.rag_enhancer.enhance_assessment_generation(
                    transcript, entities, subjective, objective
                )
                rag_context = self.rag_enhancer.format_knowledge_context(
                    enhancement.get("diagnostic_knowledge", []) + 
                    enhancement.get("reasoning_knowledge", [])
                )
            except Exception as e:
                logger.warning(f"RAG enhancement failed: {e}. Continuing without enhancement.")
        
        assessment_parts = []
        lines = transcript.split('\n')
        
        # Find diagnoses
        diagnoses = [e.text for e in entities if e.entity_type == EntityType.DIAGNOSIS]
        diseases = [e.text for e in entities if e.entity_type == EntityType.DISEASE]
        conditions = [e.text for e in entities if e.entity_type == EntityType.CONDITION]
        
        # Primary diagnosis
        all_diagnoses = list(set(diagnoses + diseases + conditions))
        if all_diagnoses:
            assessment_parts.append(f"Primary Diagnosis: {', '.join(all_diagnoses[:3])}")
        
        # Look for explicit assessment/impression statements
        assessment_keywords = ['assessment:', 'impression:', 'diagnosis:', 'clinical presentation',
                              'consistent with', 'meets criteria', 'differential diagnosis',
                              'likely', 'suggests', 'indicates', 'appears to be']
        
        assessment_statements = []
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in assessment_keywords):
                # Make sure it's a clinical assessment, not patient-reported
                if 'patient reports' not in line_lower and 'patient says' not in line_lower:
                    assessment_statements.append(line.strip())
        
        if assessment_statements:
            assessment_parts.append("\nClinical Impression:")
            assessment_parts.extend([f"  - {s}" for s in assessment_statements[:3]])
        
        # Generate differential diagnosis if multiple possibilities
        differential_keywords = ['differential', 'rule out', 'consider', 'possible', 'may be']
        differential_diagnoses = []
        
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in differential_keywords):
                differential_diagnoses.append(line.strip())
        
        if differential_diagnoses:
            assessment_parts.append("\nDifferential Diagnosis:")
            assessment_parts.extend([f"  - {d}" for d in differential_diagnoses[:3]])
        
        # Synthesize from subjective and objective if no explicit assessment
        if not assessment_parts or len(assessment_parts) < 2:
            # Combine insights from subjective symptoms and objective findings
            symptoms = [e.text for e in entities if e.entity_type == EntityType.SYMPTOM]
            vitals = [e.text for e in entities if e.entity_type == EntityType.VITAL_SIGN]
            
            clinical_reasoning = []
            if symptoms:
                clinical_reasoning.append(f"Patient presents with {', '.join(symptoms[:3])}")
            if vitals:
                clinical_reasoning.append(f"Objective findings include: {', '.join(vitals[:3])}")
            
            if clinical_reasoning:
                assessment_parts.append("\nClinical Reasoning:")
                assessment_parts.extend([f"  - {r}" for r in clinical_reasoning])
        
        return "\n".join(assessment_parts) if assessment_parts else "Assessment pending further evaluation."
    
    def _generate_plan(
        self,
        transcript: str,
        entities: List[MedicalEntity],
        assessment: str
    ) -> str:
        """
        Generate Plan section (treatment plan, medications, follow-up, goals)
        
        Plan includes (per SOAP best practices):
        - Treatment plan for next steps
        - Short-term and long-term goals
        - Specific plans for next session
        - Expectations for duration of treatment
        - Medications prescribed
        - Follow-up instructions
        - Diagnostic tests ordered
        """
        # Enhance with knowledge base if RAG is enabled
        rag_context = None
        if self.use_rag and self.rag_enhancer:
            try:
                enhancement = self.rag_enhancer.enhance_plan_generation(
                    transcript, entities, assessment
                )
                rag_context = self.rag_enhancer.format_knowledge_context(
                    enhancement.get("medication_knowledge", []) + 
                    enhancement.get("treatment_knowledge", [])
                )
            except Exception as e:
                logger.warning(f"RAG enhancement failed: {e}. Continuing without enhancement.")
        
        plan_parts = []
        lines = transcript.split('\n')
        
        # Extract medications prescribed
        medications = [e.text for e in entities if e.entity_type == EntityType.MEDICATION]
        medication_details = []
        
        for line in lines:
            line_lower = line.lower()
            if any(phrase in line_lower for phrase in ['prescribe', 'prescribed', 'medication', 
                  'medication:', 'medications:', 'take', 'mg', 'daily', 'times']):
                if 'patient' not in line_lower or 'patient reports' not in line_lower:
                    medication_details.append(line.strip())
        
        if medication_details:
            plan_parts.append("Medications:")
            plan_parts.extend([f"  - {m}" for m in medication_details[:5]])
        elif medications:
            plan_parts.append(f"Medications: {', '.join(medications)}")
        
        # Extract treatment plan
        treatment_keywords = ['treatment', 'therapy', 'intervention', 'plan:', 'treatment plan',
                             'recommend', 'advise', 'suggest', 'will continue', 'continue']
        
        treatment_plans = []
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in treatment_keywords):
                if 'patient reports' not in line_lower:
                    treatment_plans.append(line.strip())
        
        if treatment_plans:
            plan_parts.append("\nTreatment Plan:")
            plan_parts.extend([f"  - {t}" for t in treatment_plans[:4]])
        
        # Extract diagnostic tests/orders
        lab_tests = [e.text for e in entities if e.entity_type == EntityType.LAB_TEST]
        test_orders = []
        
        for line in lines:
            line_lower = line.lower()
            if any(phrase in line_lower for phrase in ['order', 'ordered', 'test', 'x-ray', 
                  'chest x-ray', 'lab', 'diagnostic', 'imaging']):
                if 'patient reports' not in line_lower:
                    test_orders.append(line.strip())
        
        if test_orders:
            plan_parts.append("\nDiagnostic Tests/Orders:")
            plan_parts.extend([f"  - {t}" for t in test_orders[:4]])
        elif lab_tests:
            plan_parts.append(f"\nDiagnostic Tests Ordered: {', '.join(lab_tests)}")
        
        # Extract follow-up instructions
        followup_keywords = ['follow-up', 'follow up', 'followup', 'return', 'next appointment',
                            'scheduled', 'appointment', 'revisit', 'come back']
        
        followup_instructions = []
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in followup_keywords):
                followup_instructions.append(line.strip())
        
        if followup_instructions:
            plan_parts.append("\nFollow-up Instructions:")
            plan_parts.extend([f"  - {f}" for f in followup_instructions[:3]])
        
        # Extract goals (short-term and long-term)
        goal_keywords = ['goal', 'goals', 'short-term', 'long-term', 'objective', 'target',
                        'expect', 'expectation', 'duration', 'outcome']
        
        goals = []
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in goal_keywords):
                goals.append(line.strip())
        
        if goals:
            plan_parts.append("\nGoals:")
            plan_parts.extend([f"  - {g}" for g in goals[:3]])
        
        # Extract patient instructions/education
        instruction_keywords = ['instruct', 'advise', 'educate', 'counsel', 'recommend that',
                               'patient should', 'advised to', 'instructed to']
        
        instructions = []
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in instruction_keywords):
                instructions.append(line.strip())
        
        if instructions:
            plan_parts.append("\nPatient Instructions/Education:")
            plan_parts.extend([f"  - {i}" for i in instructions[:3]])
        
        # If no structured extraction, create basic plan
        if not plan_parts or len(plan_parts) < 2:
            if medications:
                plan_parts.append(f"Continue current medications: {', '.join(medications)}")
            plan_parts.append("Treatment plan to be determined based on assessment.")
            plan_parts.append("Follow-up appointment scheduled as needed.")
        
        return "\n".join(plan_parts) if plan_parts else "Plan pending assessment completion."

